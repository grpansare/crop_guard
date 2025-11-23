#!/usr/bin/env python3
"""
Convert PROJECT_DOCUMENTATION.md to Word (.docx) format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def process_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and not line.startswith('|'):
            # End of table, process it
            if len(table_lines) > 2:
                # Parse table
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                num_cols = len(headers)
                num_rows = len(table_lines) - 1  # Exclude separator line
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                for j, header in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                
                # Add data rows (skip separator line at index 1)
                row_idx = 1
                for table_line in table_lines[2:]:
                    cells_data = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    for j, cell_data in enumerate(cells_data):
                        if j < num_cols and row_idx < num_rows:
                            table.rows[row_idx].cells[j].text = cell_data
                    row_idx += 1
                
                doc.add_paragraph()  # Add spacing after table
            
            in_table = False
            table_lines = []
            # Don't increment i, process current line
            continue
        
        # Skip horizontal rules
        if line.strip() == '---' or line.strip() == '___':
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level <= 6:
                add_heading(doc, text, level=min(level, 3))
            i += 1
            continue
        
        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Check for checkboxes
            if text.startswith('[x]'):
                text = '✓ ' + text[4:]
            elif text.startswith('[/]'):
                text = '◐ ' + text[4:]
            elif text.startswith('[ ]'):
                text = '☐ ' + text[4:]
            
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Handle bold/italic text in paragraphs
        if line.strip():
            # Simple bold (**text**)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            # Remove markdown links but keep text
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            # Remove inline code backticks
            text = re.sub(r'`(.+?)`', r'\1', text)
            
            if '**' in line or '*' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(text)
        else:
            # Empty line
            if i > 0 and lines[i-1].strip():  # Only add if previous line had content
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    process_markdown_to_docx("PROJECT_DOCUMENTATION.md", "PROJECT_DOCUMENTATION.docx")
